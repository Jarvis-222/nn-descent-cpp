#!/usr/bin/env python3
"""Generate the workshop paper draft as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper functions ──

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_status_tag(paragraph, status):
    """Append a colored status tag to a paragraph."""
    run = paragraph.add_run(f"  [{status}]")
    if status in ("DONE", "READY", "EXISTS"):
        run.font.color.rgb = RGBColor(0, 128, 0)
    elif status in ("NEED", "TODO", "NEED TO BUILD", "PENDING"):
        run.font.color.rgb = RGBColor(200, 0, 0)
    elif status == "OPTIONAL":
        run.font.color.rgb = RGBColor(180, 120, 0)
    run.bold = True
    run.font.size = Pt(9)

def add_para(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.3 + 0.3 * level)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_checklist(items):
    """items: list of (text, done_bool)"""
    for text, done in items:
        sym = "\u2611" if done else "\u2610"
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(f"{sym} {text}")
        if done:
            run.font.color.rgb = RGBColor(0, 128, 0)
        else:
            run.font.color.rgb = RGBColor(200, 0, 0)
        run.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Workshop Paper Draft\n")
run.bold = True
run.font.size = Pt(22)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "Projection-Based Distance Filtering for\n"
    "Accelerating NN-Descent Graph Construction"
)
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("\nJainish Pandya\n").bold = True
meta.add_run("M.Eng. Electrical & Computer Engineering\n")
meta.add_run("University of Waterloo\n\n")
meta.add_run("Target: ").bold = True
meta.add_run("8 pages + references (IEEE two-column)\n")
meta.add_run("Status: ").bold = True
meta.add_run("DRAFT — Structure & Gap Analysis\n")
meta.add_run("Date: ").bold = True
meta.add_run("March 2026\n")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PAGE BUDGET
# ══════════════════════════════════════════════════════════════

add_heading("Page Budget Overview", level=1)

add_table(
    ["Section", "Pages", "Status"],
    [
        ["Abstract", "0.25", "TO WRITE"],
        ["1. Introduction", "1.25", "TO WRITE (material ready from survey)"],
        ["2. Related Work & Background", "1.5", "TO WRITE (can adapt from survey)"],
        ["3. Method", "1.5", "TO WRITE (code + math exists)"],
        ["4. Experiments & Results", "2.5", "PARTIALLY DONE"],
        ["5. Discussion", "0.5", "TO WRITE"],
        ["6. Conclusion & Future Work", "0.5", "TO WRITE"],
        ["References", "~1", "MOSTLY READY (from survey)"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 1: INTRODUCTION
# ══════════════════════════════════════════════════════════════

add_heading("1. Introduction (~1.25 pages)", level=1)

p = add_para("Story Arc", bold=True)
add_para(
    "ANN is critical infrastructure \u2192 field evolved through hashing / trees / "
    "quantization / graphs \u2192 now hybrid era \u2192 NN-Descent still matters "
    "(CAGRA, parallelizability) \u2192 its bottleneck is wasted distance computation "
    "\u2192 we propose projection-based filtering to fix it.",
    italic=True,
    indent=True,
)

doc.add_paragraph()
add_heading("Paragraph-by-Paragraph Plan", level=2)

# Para 1
p = add_para("Para 1 \u2014 ANN matters now more than ever", bold=True)
add_bullet("LLMs, RAG, vector databases \u2192 ANN is infrastructure")
add_bullet("Exact search is O(nd), prohibitive at scale")
add_bullet("Source: Adapt directly from survey paper Section I")
add_status_tag(p, "READY")

# Para 2
p = add_para("Para 2 \u2014 Evolutionary view (CONDENSED to one paragraph)", bold=True)
add_bullet("Hashing (LSH, 1999) \u2192 Trees (KD-tree, RP-tree) \u2192 Quantization (PQ, 2011) \u2192 Graphs (NSW \u2192 HNSW)")
add_bullet('Key insight: field moved from "search-only" structures to "build expensive, search fast"')
add_bullet("Timeline: LSH (1999) \u2192 PQ (2011) \u2192 NN-Descent (2011) \u2192 HNSW (2018/2020)")
add_bullet("Source: Material exists in survey Sections III\u2013IV. CONDENSE to ~6\u20138 sentences.")
add_status_tag(p, "READY")

# Para 3
p = add_para("Para 3 \u2014 Hybrid era: combining strengths", bold=True)
add_bullet("PyNNDescent = NN-Descent + RP-trees (init)")
add_bullet("LSH-APG = LSH + proximity graph construction")
add_bullet("FINGER = low-rank projections + HNSW search")
add_bullet("SHG = shortcuts + compressed upper layers")
add_bullet("Pattern: lightweight structures used to FILTER or ACCELERATE graph operations")
add_status_tag(p, "READY")

# Para 4
p = add_para("Para 4 \u2014 Why NN-Descent still matters", bold=True)
add_bullet("HNSW dominates search, but is CPU-bound and hard to parallelize")
add_bullet("NN-Descent is backbone of CAGRA (NVIDIA GPU graph construction)")
add_bullet("Also used in EFANNA, PyNNDescent")
add_bullet("Simple, embarrassingly parallel local-join structure")
add_bullet("Yet: main bottleneck is WASTED distance computations during construction")
add_status_tag(p, "NEED")
add_para("\u26a0 Need: CAGRA citation, verify HNSW parallelization limitation claim.", italic=True, indent=True)

# Para 5
p = add_para("Para 5 \u2014 The waste problem + our contribution", bold=True)
add_bullet("[FIGURE 1: Waste ratio \u2014 fraction of dist_comps that don't update graph per iteration]")
add_bullet("Motivates: can we SKIP these wasted computations cheaply?")
add_bullet("Contribution: adapt projection-based filtering (from LSH-APG) to NN-Descent")
add_bullet("First application of this filtering to NN-Descent")
add_bullet("1.1\u20131.4x wall-clock speedup on high-dim data, <3% recall loss at conservative settings")
add_status_tag(p, "NEED")
add_para("\u26a0 Need: Build the waste-ratio figure from existing iteration stats.", italic=True, indent=True)

add_heading("Checklist for Introduction", level=2)
add_checklist([
    ("Evolutionary narrative (material from survey)", True),
    ("Hybrid algorithms discussion (survey Section V)", True),
    ("Contribution statement", True),
    ('"Waste ratio" figure (build from iteration stats)', False),
    ("CAGRA citation / verification it uses NN-Descent", False),
    ("HNSW parallelization limitations claim + citation", False),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 2: RELATED WORK & BACKGROUND
# ══════════════════════════════════════════════════════════════

add_heading("2. Related Work & Background (~1.5 pages)", level=1)

add_heading("2.1 NN-Descent Algorithm (~0.5 page)", level=2)
add_bullet("Algorithm overview: random init \u2192 local join \u2192 new/old partition \u2192 convergence")
add_bullet("Complexity: O(n^1.14) empirically (Dong et al. 2011)")
add_bullet("Key parameters: k (neighbors), \u03c1 (sampling), \u03b4 (convergence)")
add_bullet('The "local join" is where all distance computations happen')
p = add_para(""); add_status_tag(p, "READY")

add_heading("2.2 Filtering in Graph-Based ANN (~0.5 page)", level=2)
add_bullet("FINGER (Chen et al. 2023): low-rank projections for HNSW search \u2014 20\u201360% speedup")
add_bullet("LSH-APG (Zhang et al. 2023): LSH for graph construction + search entry points")
add_bullet("    \u2192 They use projection-based filtering during HNSW construction")
add_bullet("    \u2192 Our work adapts this to NN-Descent's local-join structure")
add_bullet("SHG (Gong et al. 2025): shortcuts + compression for HNSW search")
p = add_para(""); add_status_tag(p, "NEED")
add_para("\u26a0 Need: Re-read LSH-APG paper to precisely describe their filter vs. ours.", italic=True, indent=True)

add_heading("2.3 NN-Descent Variants (~0.5 page)", level=2)
add_bullet("PyNNDescent: RP-tree init + optimized local join (Numba JIT)")
add_bullet("EFANNA: NN-Descent + KD-tree init + graph diversification")
add_bullet("CAGRA (NVIDIA): GPU-accelerated NN-Descent for graph construction")
add_bullet("Key point: NONE of these use distance filtering during local join. That's our gap.")
p = add_para(""); add_status_tag(p, "NEED")
add_para("\u26a0 Need: Verify CAGRA/EFANNA details. Verify no prior filtering work on NN-Descent.", italic=True, indent=True)

add_checklist([
    ("All references from survey paper", True),
    ("FINGER, LSH-APG, SHG descriptions", True),
    ("NN-Descent algorithm understanding", True),
    ("Re-read LSH-APG paper for precise comparison", False),
    ("Verify CAGRA uses NN-Descent", False),
    ("Confirm no prior work applies filtering to NN-Descent", False),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 3: METHOD
# ══════════════════════════════════════════════════════════════

add_heading("3. Method (~1.5 pages)", level=1)

add_heading("3.1 Collision-Count Filter \u2014 Negative Result (~0.25 page)", level=2)
add_bullet("Idea: LSH fingerprints (128-bit) \u2192 Hamming distance as similarity proxy")
add_bullet("Filter rule: skip pair if collision_count < threshold")
add_bullet("Why it failed: discrete, noisy, poorly calibrated")
add_bullet("Show brief table of collision filter results vs. no filter")
add_bullet("Purpose: motivates why we need a theoretically-grounded approach")
p = add_para(""); add_status_tag(p, "NEED")
add_para("\u26a0 Need: Collision filter experimental results (user to provide).", italic=True, indent=True)

add_heading("3.2 Projection-Based Distance Filter (~0.75 page)", level=2)
add_para("Mathematical Foundation:", bold=True)
add_para(
    "Given m random Gaussian vectors a\u2081, ..., a\u2098, define projections:\n"
    "    P(x) = (a\u2081\u00b7x, ..., a\u2098\u00b7x)\n\n"
    "For two points u, v with true L2 distance d(u,v):\n"
    "    ||P(u) - P(v)||\u00b2 / d(u,v)\u00b2  ~  \u03c7\u00b2(m) / m\n\n"
    "By chi-squared concentration:\n"
    "    P( ||P(u)-P(v)||\u00b2 > t\u00b2 \u00b7 d(u,v)\u00b2 ) < 1 - p\u03c4\n\n"
    "where t\u00b2 = \u03c7\u00b2_p\u03c4(m) (inverse CDF of chi-squared with m degrees of freedom).",
    indent=True,
)

add_para("Filter Condition:", bold=True)
add_para(
    "Skip pair (u1, u2) if:\n"
    "    proj_dist_sq(u1, u2) > t\u00b2 \u00b7 dk(u1)\u00b2   AND\n"
    "    proj_dist_sq(u1, u2) > t\u00b2 \u00b7 dk(u2)\u00b2\n\n"
    "where dk(u) = distance to u's current k-th nearest neighbor.\n"
    "Conservative AND logic: both points must agree the pair is far.",
    indent=True,
)

add_para("Parameters:", bold=True)
add_bullet("m = number of projections (default: 32)")
add_bullet("p\u03c4 = confidence level (default: 0.95)")
p = add_para(""); add_status_tag(p, "READY")

add_heading("3.3 Modified NN-Descent Algorithm (~0.25 page)", level=2)
add_bullet("Pseudocode box showing where filter plugs into local-join loop")
add_bullet("One-time setup: O(n \u00b7 m \u00b7 d) to compute all projections")
add_bullet("Per-pair filter cost: O(m) vs. O(d) for true distance")
add_bullet("When d >> m: huge savings. When d \u2248 m: overhead \u2248 savings.")
p = add_para(""); add_status_tag(p, "READY")
add_para("\u26a0 Need: Create clean pseudocode figure for the paper.", italic=True, indent=True)

add_heading("3.4 Complexity Analysis (~0.25 page)", level=2)
add_bullet("Memory: O(n \u00b7 m) additional for projection storage")
add_bullet("Setup: O(n \u00b7 m \u00b7 d) one-time projection computation")
add_bullet("Per iteration: each pair costs O(m) for filter check instead of O(d)")
add_bullet("If filter rate = f, effective per-pair cost = (1\u2212f)\u00b7O(d) + O(m)")
add_bullet("Net savings when f \u00b7 d > m (approximately)")
p = add_para(""); add_status_tag(p, "READY")

add_heading("3.5 Cosine Distance Extension", level=2)
add_para(
    "L2: chi-squared bound is well-established via Johnson-Lindenstrauss lemma.\n"
    "Cosine: need to verify concentration bound applies. May require different "
    "inequality (e.g., Charikar 2002 random hyperplane analysis).\n"
    "Decision: include in paper or defer to future work?",
    italic=True,
    indent=True,
)
p = add_para(""); add_status_tag(p, "PENDING")

add_checklist([
    ("Projection filter math (implemented + tested)", True),
    ("Algorithm pseudocode (extractable from code)", True),
    ("Complexity analysis", True),
    ("Collision filter experimental results", False),
    ("Clean pseudocode figure", False),
    ("Cosine distance math verification", False),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 4: EXPERIMENTS & RESULTS
# ══════════════════════════════════════════════════════════════

add_heading("4. Experiments & Results (~2.5 pages)", level=1)

add_heading("4.1 Experimental Setup (~0.25 page)", level=2)
add_bullet("Datasets: GIST 1M (960-d), GIST 100K (960-d), SIFT 1M (128-d)")
add_bullet("Ground truth: brute-force via FAISS GPU")
add_bullet("Metrics: Recall@10, distance computations, wall-clock time")
add_bullet("Parameters: k=10, mc=40, m=32 projections, random init")
p = add_para(""); add_status_tag(p, "NEED")
add_para("\u26a0 Need: Document hardware specs (CPU model, RAM, compiler flags).", italic=True, indent=True)

# --- 4.2 Main results ---
add_heading("4.2 Main Result: Projection Filter Effectiveness (~0.75 page)", level=2)

add_para("Table: GIST 1M \u2014 Projection Filter Sweep (mc=40, m=32)", bold=True)
add_table(
    ["Config", "Recall", "Dist Comps", "Time (s)", "Speedup", "Recall Loss"],
    [
        ["No filter",   "0.5894", "1,551.5M", "1056.7", "1.00x", "0.00%"],
        ["p\u03c4 = 0.99", "0.5879", "1,425.5M", "1058.1", "1.00x", "0.25%"],
        ["p\u03c4 = 0.95", "0.5820", "1,265.7M", "962.1",  "1.10x", "1.26%"],
        ["p\u03c4 = 0.90", "0.5720", "1,143.0M", "885.8",  "1.19x", "2.95%"],
        ["p\u03c4 = 0.85", "0.5603", "1,048.4M", "829.0",  "1.27x", "4.94%"],
        ["p\u03c4 = 0.80", "0.5470", "968.3M",   "778.5",  "1.36x", "7.19%"],
    ],
)
p = add_para(""); add_status_tag(p, "DONE")
add_para("Also have: GIST 100K (5-run averaged), SIFT 1M (5-run averaged).", italic=True, indent=True)

# --- 4.3 Filter vs Sampling ---
add_heading("4.3 Filter vs. Sampling Reduction (~0.5 page)", level=2)
add_para(
    "Key argument: reducing max_candidates (mc) also reduces dist_comps, "
    "but filtering gives BETTER recall per dist_comp.",
    indent=True,
)

add_para("Table: GIST 1M \u2014 MC Sweep, No Filter", bold=True)
add_table(
    ["mc", "Recall", "Dist Comps", "Time (s)"],
    [
        ["25", "0.5207", "1,250.9M", "871.8"],
        ["30", "0.5490", "1,364.9M", "942.0"],
        ["35", "0.5713", "1,463.7M", "1003.6"],
        ["40", "0.5894", "1,551.5M", "1056.7"],
    ],
)
p = add_para(""); add_status_tag(p, "DONE")
add_para("Plot: overlay filter sweep + mc sweep on same recall-vs-distcomp axes.", italic=True, indent=True)

# --- 4.4 Dimensionality ---
add_heading("4.4 Dimensionality Analysis (~0.5 page)", level=2)
add_bullet("GIST (960-d): filter O(32) << distance O(960) \u2192 significant speedup (1.1\u20131.4x)")
add_bullet("SIFT (128-d): filter O(32) \u2248 distance O(128) \u2192 marginal speedup (~1.01\u20131.09x)")
add_bullet("This EXPLAINS when to use the filter: high-dimensional data where d >> m")
add_bullet("Cross-dataset bar charts for speedup and recall loss already exist")
p = add_para(""); add_status_tag(p, "DONE")

# --- 4.5 Generality ---
add_heading("4.5 Generality Across Initializers (~0.25 page)", level=2)
add_bullet("Show filter works with random init, LSH init, RP-tree init")
add_bullet("Small table (3 rows) showing comparable speedup across init methods")
add_bullet("Demonstrates this is a general NN-Descent enhancement")
p = add_para(""); add_status_tag(p, "NEED")
add_para("\u26a0 Need: Run experiments on GIST 100K with LSH init + filter and RP-tree init + filter.", italic=True, indent=True)

# --- 4.6 Optional ---
add_heading("4.6 [OPTIONAL] Cosine Distance Results (~0.25 page)", level=2)
add_para("If math checks out: run filter on GIST 100K with cosine distance.", italic=True, indent=True)
p = add_para(""); add_status_tag(p, "PENDING")

add_heading("4.7 [OPTIONAL] Collision Filter Results (~0.25 page)", level=2)
add_para("Brief table showing collision filter didn't help. Motivates projection approach.", italic=True, indent=True)
add_para("Could go in Method section instead.", italic=True, indent=True)
p = add_para(""); add_status_tag(p, "NEED")

add_heading("Experiments Checklist", level=2)
add_checklist([
    ("GIST 1M: projection filter sweep (single run)", True),
    ("GIST 100K: projection filter sweep (5-run averaged)", True),
    ("SIFT 1M: projection filter sweep (5-run averaged)", True),
    ("GIST 1M: MC sweep (sampling baseline)", True),
    ("GIST 100K: MC sweep", True),
    ("All plots via plot_all.py", True),
    ("Collision filter results (user providing)", False),
    ("Different initializer + filter runs (GIST 100K)", False),
    ("Hardware specification", False),
    ("[OPTIONAL] Cosine distance experiments", False),
    ("[OPTIONAL] PyNNDescent comparison", False),
    ("Paper-quality plot styling", False),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 5: DISCUSSION
# ══════════════════════════════════════════════════════════════

add_heading("5. Discussion (~0.5 page)", level=1)

add_para("Key Points to Make:", bold=True)
doc.add_paragraph()

add_para("1. When it helps:", bold=True)
add_para(
    "High-dimensional data where d >> m. Filter check O(m) is cheap relative to true "
    "distance O(d). GIST (960-d) shows clear benefit.",
    indent=True,
)

add_para("2. When it doesn't:", bold=True)
add_para(
    "Low-dimensional data (SIFT 128-d). Filter overhead \u2248 savings. "
    "Recommendation: don't use for d < ~200.",
    indent=True,
)

add_para("3. Construction-time recall vs. query-time quality:", bold=True)
add_para(
    "Recall loss is at graph construction. For downstream tasks (HNSW search, CAGRA queries), "
    "a graph with 0.58 vs 0.59 construction recall may perform identically at query time. "
    "Needs investigation but is a strong practical argument.",
    indent=True,
)

add_para("4. Applicability:", bold=True)
add_para(
    "This is a DROP-IN enhancement. PyNNDescent, EFANNA, CAGRA can all add projection "
    "filtering to their local-join loop. No structural changes needed.",
    indent=True,
)

add_para("5. Memory overhead:", bold=True)
add_para(
    "O(n \u00b7 m) floats for projections. For n=1M, m=32: ~128 MB. Modest.",
    indent=True,
)

p = add_para(""); add_status_tag(p, "READY")

# ══════════════════════════════════════════════════════════════
# SECTION 6: CONCLUSION & FUTURE WORK
# ══════════════════════════════════════════════════════════════

add_heading("6. Conclusion & Future Work (~0.5 page)", level=1)

add_para("Conclusion:", bold=True)
add_bullet("Projection-based filtering reduces distance computations in NN-Descent by 18\u201337%")
add_bullet("Wall-clock speedup of 1.1\u20131.4x on high-dimensional data")
add_bullet("Theoretically grounded (chi-squared), tunable (p\u03c4), conservative (AND logic)")
add_bullet("Drop-in enhancement for any NN-Descent variant")

add_para("Future Work:", bold=True)
add_bullet("Cosine / inner-product distance: Extend filter to non-L2 metrics")
add_bullet("GPU integration: Apply to CAGRA's GPU NN-Descent (projection check is SIMD-friendly)")
add_bullet("Adaptive p\u03c4: Start conservative (0.99) early, increase aggressiveness as graph stabilizes")
add_bullet("Learned projections: Replace random Gaussian with data-dependent projections (PCA)")
add_bullet("Query-time impact study: Measure how construction recall loss affects downstream search")

p = add_para(""); add_status_tag(p, "READY")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# FIGURES INVENTORY
# ══════════════════════════════════════════════════════════════

add_heading("Figures Inventory", level=1)

add_table(
    ["Figure", "Description", "Status"],
    [
        ["Fig 1", "Waste ratio: % of dist_comps not updating graph per iteration", "NEED TO BUILD"],
        ["Fig 2", "GIST 1M: Recall vs Dist Comps (filter + mc sweep overlay)", "EXISTS"],
        ["Fig 3", "GIST 1M: Recall vs Wall-Clock Time", "EXISTS"],
        ["Fig 4", "Cross-dataset speedup bar chart", "EXISTS"],
        ["Fig 5", "Cross-dataset recall loss bar chart", "EXISTS"],
        ["Fig 6", "Distance computation savings (%) bar chart", "EXISTS"],
        ["Fig 7", "[OPT] Per-iteration filter rate evolution", "COULD BUILD"],
        ["Fig 8", "Algorithm pseudocode box", "NEED TO CREATE"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# OPEN DECISIONS
# ══════════════════════════════════════════════════════════════

add_heading("Open Decisions for Professor", level=1)

add_para("1. PyNNDescent Comparison", bold=True)
add_bullet("Option A: Compare wall-clock (risk: PyNNDescent is faster due to optimized NumPy/Numba)")
add_bullet("Option B: Compare distance computations only \u2014 frame as algorithmic improvement")
add_bullet("Option C: Frame as complementary enhancement (PyNNDescent could adopt this too)")
add_para("Recommendation: Option B or C (safer, avoids implementation fairness debate).", italic=True, indent=True)

add_para("2. Cosine Distance", bold=True)
add_bullet("Include: strengthens paper (cosine is more expensive per call \u2192 bigger filter payoff)")
add_bullet("Defer: math needs verification (chi-squared bound may not directly apply)")
add_para("Recommendation: Verify math first, include if straightforward.", italic=True, indent=True)

add_para("3. Paper Framing", bold=True)
add_bullet('Option A: "Enhancement to NN-Descent" (narrower, cleaner)')
add_bullet('Option B: "Enhancement to any local-join graph construction" (broader, needs more evidence)')
add_para("Recommendation: Option A for now, mention broader applicability in Discussion.", italic=True, indent=True)

add_para("4. Collision Filter", bold=True)
add_bullet("Include as negative result in Method section (strengthens narrative)")
add_bullet("Or just mention briefly in one sentence")
add_para("Recommendation: Include as 0.25-page subsection \u2014 shows we explored alternatives.", italic=True, indent=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ACTION ITEMS
# ══════════════════════════════════════════════════════════════

add_heading("Immediate Action Items (Priority Order)", level=1)

add_table(
    ["#", "Task", "Effort", "Blocking?"],
    [
        ["1", "Share collision filter results", "User action", "Yes (Method section)"],
        ["2", "Run filter with LSH + RP-tree init on GIST 100K", "~2h compute", "Yes (generality claim)"],
        ["3", 'Build "waste ratio" plot from iteration stats', "~1h coding", "Yes (intro Figure 1)"],
        ["4", "Document hardware specs", "5 min", "Yes (experiments)"],
        ["5", "Verify cosine distance math", "~2h research", "Blocks cosine decision"],
        ["6", "Re-read LSH-APG paper", "~1h", "Blocks related work"],
        ["7", "Verify CAGRA uses NN-Descent + get citation", "~30 min", "Blocks intro"],
        ["8", "Regenerate plots with paper-quality styling", "~2h", "No (can do last)"],
        ["9", "Decide PyNNDescent framing with professor", "Discussion", "Blocks experiment scope"],
    ],
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════

add_heading("References Needed", level=1)

add_para("Already have (from survey):", bold=True)
add_bullet("[Dong et al. 2011] NN-Descent")
add_bullet("[Malkov & Yashunin 2020] HNSW")
add_bullet("[Zhang et al. 2023] LSH-APG")
add_bullet("[Chen et al. 2023] FINGER")
add_bullet("[Gong et al. 2025] SHG")
add_bullet("[J\u00e9gou et al. 2011] PQ / Faiss")
add_bullet("[ANN-Benchmarks]")

add_para("Need to add:", bold=True)
add_checklist([
    ("CAGRA paper (NVIDIA, Ootomo et al. 2023 or similar)", False),
    ("EFANNA paper (Fu & Cai, 2016)", False),
    ("PyNNDescent (McInnes, 2018)", False),
    ("Johnson-Lindenstrauss lemma reference", False),
    ("Chi-squared concentration bounds reference", False),
])

# ── Save ──
out = "/home/jainish/UWaterloo/Winter_2026/nn-descent-cpp/Workshop_Paper_Draft.docx"
doc.save(out)
print(f"Saved: {out}")
